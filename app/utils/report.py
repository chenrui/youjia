#!/usr/bin/env python
# -*- coding: utf-8 -*-

import os
import io
import sys
import json
import uuid
import shutil
import zipfile
import requests
import subprocess32

from string import Formatter
from docx import Document


#
# Define global constants
#

# all paragraph definition

P0_IDX_MAIN_TITLE                            = 0
P1_IDX_MAIN_SUBTITLE                         = 1

P2_IDX_SUMMARY_TITLE                         = 2
P3_IDX_SUMMARY_CONTENT                       = 3

P4_IDX_STATISTICS_DATA                       = 4
P5_IDX_STATISTICS_TEXT                       = 5

P7_IDX_DEFENSE_SUBJECT                       = 7

P8_IDX_DEFENSE_ATTACK_TREND_TITLE            = 8
P9_IDX_DEFENSE_ATTACK_TREND_CONTENT          = 9
P10_IDX_DEFENSE_ATTACK_TREND_PICTURE         = 10

P11_IDX_DEFENSE_ATTACK_WEBSITE_TITLE         = 11
P12_IDX_DEFENSE_ATTACK_WEBSITE_CONTENT       = 12
P13_IDX_DEFENSE_ATTACK_WEBSITE_PICTURE       = 13

P14_IDX_DEFENSE_ATTACK_DISTRIBUTION_TITLE    = 14
P15_IDX_DEFENSE_ATTACK_DISTRIBUTION_CONTENT  = 15
P16_IDX_DEFENSE_ATTACK_DISTRIBUTION_PICTURE = 16

P17_IDX_DEFENSE_SOURCE_IPADDRESS_TITLE       = 17
P18_IDX_DEFENSE_SOURCE_IPADDRESS_CONTENT     = 18
P19_IDX_DEFENSE_SOURCE_IPADDRESS_PICTURE     = 19

P20_IDX_DEFENSE_SOURCE_REGION_TITLE          = 20
P21_IDX_DEFENSE_SOURCE_REGION_CONTENT        = 21
P22_IDX_DEFENSE_SOURCE_REGION_PICTURE        = 22

P23_IDX_ACCELERATION_SUBJECT                 = 23

P24_IDX_ACCELERATION_REQUESTS_TITLE          = 24
P25_IDX_ACCELERATION_REQUESTS_CONTENT        = 25
P26_IDX_ACCELERATION_REQUESTS_PICTURE        = 26

P27_IDX_ACCELERATION_TRAFFIC_TITLE           = 27
P28_IDX_ACCELERATION_TRAFFIC_CONTENT         = 28
P29_IDX_ACCELERATION_TRAFFIC_PICTURE         = 29

P30_IDX_ACCELERATION_IPADDRESS_TITLE         = 30
P31_IDX_ACCELERATION_IPADDRESS_CONTENT       = 31
P32_IDX_ACCELERATION_IPADDRESS_PICTURE       = 32

P33_IDX_ACCELERATION_PV_TITLE                = 33
P34_IDX_ACCELERATION_PV_CONTENT              = 34
P35_IDX_ACCELERATION_PV_PICTURE              = 35

P36_IDX_ACCELERATION_TREND_TITLE             = 36
P37_IDX_ACCELERATION_TREND_PICTURE           = 37

P38_IDX_ACCELERATION_REGION_TITLE            = 38
P39_IDX_ACCELERATION_REGION_PICTURE          = 39

P99_IDX_MAXIMUM_VALUE                        = 40


# all paragraph run definition

R_IDX_P0_TITLE              = 0

R_IDX_P1_RANGE_TIME         = 1
R_IDX_P1_EXPORT_TIME        = 4

R_IDX_P3_RANGE_TIME         = 1
R_IDX_P3_SUBDOMAIN          = 3
R_IDX_P3_TOTAL_ATTACK       = 7
R_IDX_P3_WEB_ATTACK         = 11
R_IDX_P3_CC_ATTACK          = 15
R_IDX_P3_IPADDRESS          = 21
R_IDX_P3_TOTAL_REQUESTS     = 27
R_IDX_P3_ACC_TIMES          = 32
R_IDX_P3_TOTAL_TRAFFIC      = 37
R_IDX_P3_SAVED_TRAFFIC      = 42

R_IDX_P4_T1_S1              = 6
R_IDX_P4_T1_S2              = 7
R_IDX_P4_T2_S1              = 10
R_IDX_P4_T2_S2              = 11
R_IDX_P4_T3_S1              = 13
R_IDX_P4_T3_S2              = 14
R_IDX_P4_T4_S1              = 16
R_IDX_P4_T4_S2              = 17
R_IDX_P4_T5_S1              = 20
R_IDX_P4_T5_S2              = 21
R_IDX_P4_T6_S1              = 24
R_IDX_P4_T6_S2              = 25

R_IDX_P9_RANGE_TIME         = 2
R_IDX_P9_LOW_TIME           = 5
R_IDX_P9_LOW_ATTACK_TIMES   = 9
R_IDX_P9_HIGH_TIME          = 13
R_IDX_P9_HIGH_ATTACK_TIMES  = 17

R_IDX_P12_RANGE_TIME        = 2
R_IDX_P12_TOTAL_DOMAIN      = 6
R_IDX_P12_TOP_DOMAIN        = 10
R_IDX_P12_TOP_COUNT         = 13

R_IDX_P15_RANGE_TIME        = 1
R_IDX_P15_ATTACK_CATEGORY   = 3
R_IDX_P15_ATTACK_TIMES      = 5
R_IDX_P15_PERCENTAGE        = 7

R_IDX_P18_RANGE_TIME        = 1
R_IDX_P18_IP_ADDRESS        = 4
R_IDX_P18_ATTACK_TIMES      = 6

R_IDX_P21_RANGE_TIME        = 1
R_IDX_P21_REGION            = 3
R_IDX_P21_ATTACK_TIMES      = 6

R_IDX_P25_RANGE_TIME        = 1
R_IDX_P25_TOTAL_REQUESTS    = 3
R_IDX_P25_ACC_REQUESTS      = 6

R_IDX_P28_RANGE_TIME        = 1
R_IDX_P28_TOTAL_TRAFFIC     = 4
R_IDX_P28_SAVED_TRAFFIC     = 7

R_IDX_P31_RANGE_TIME        = 1
R_IDX_P31_TOTAL_IP          = 4
R_IDX_P31_TOP_TIME          = 6
R_IDX_P31_TOP_IP            = 8

R_IDX_P34_RANGE_TIME        = 1
R_IDX_P34_TOTAL_PV          = 3
R_IDX_P34_TOP_TIME          = 5
R_IDX_P34_TOP_PV            = 7


# TX display number width definition
#
# ----------------------- 
# -  1000  - 百万次
# - NUMBER - UNIT FIELD
# -----------------------

# total A4 width measured by space width
TX_TOTAL_WIDTH      = 230

# each TX total width
TX_EACH_WIDTH       = 30

# one-byte char width in unit field (space width)
TX_UNIT_BASIC_WIDTH = 1.25

# two-byte char width (Chinese Chars)
TX_UNIT_DWORD_WIDTH = 2.5

# one-byte char width in number field
TX_NUMBER_WIDTH     = 5

# one-byte sign char width in number field
TX_SIGN_WIDTH       = 7

# two-byte char width (Chinese Chars) in number field
TX_DWORD_WITDTH     = 8.75



#
# Define global variables
#

embedded_placeholder_sha1 = [
    '8a12facd63da097a924893b30222ca59848038bf',
    'b7aff7007b2b2119ce24bdabecb62092565d609c',
    '8f0dbdf24d5e784d9346096271ce9605943cb840',
    '1a9755766c1e8265acbdc38df2f59078265a274c',
    '61ddc89d11fecb6ec201db916e8c6b280d03a167',
    'fe37e3429123cdef76a811835891ae4b2e11bede',
    '4975d85a5997e6fd9dcd71c2adae6b586090772a',
    '3ee8e2500533e308734e5981e3776b7c18442d33',
    '02102d8956d1407a545242915ff2d4394158ea8c',
    'd055fa505de71214d072996aa909c34eca83f95b',
    '0dfb986e9c4a586256731281b34d24b8199cace1'
]

workspace_home      = os.path.abspath(os.path.dirname(__file__))

template_v1_h1_path = os.path.join(workspace_home, '..', 'templates', 'report', 'template_v1_h1.docx')


class DocxFormat(object):

    def __init__(self, data, placeholders, path):
        self._data = data
        self._doc = None
        self._docx = None
        self._path = path
        self._range_time = None
        self._formatter = Formatter()
        self._img_index = 0
        self._sha1_index = 0
        self._img_prefix = 'report-'
        self._img_suffix = '.jpg'
        self._img_placeholders = placeholders
        self.__p_index_array = []

    def format_datetime(self, datetime, hyphen = False):
        if hyphen:
            dt = datetime.split('-')
            return dt[0] + u'年' + dt[1] + u'月' + dt[2] + u'日'
        else:
            return datetime[0:4] + u'年' + datetime[4:6] + u'月' + datetime[6:8] + u'日'

    def _format_number_unit(self, value):
        if value <= 9999:
            width = 0
            unit  = u""
            mul   = value
        elif value < 99999999:
            width = 1
            unit  = u"万"
            mul   = value/10000
        elif value < 999999999999:
            width = 1
            unit  = u"亿"
            mul   = value/100000000
        elif value < 9999999999999999:
            width = 2
            unit  = u"万亿"
            mul   = value/1000000000000
        else:
            return (u"1", u">万万亿", 1, 3.5)

        return (int(mul), unit, len(str(int(mul))), width)

    def _format_traffic_unit(self, value):
        if value <= 1023:
            width = 1
            unit  = u"B"
            mul   = value
        elif value < 1048575:
            width = 2
            unit  = u"KB"
            mul   = value/1024
        elif value < 1073741823:
            width = 2
            unit  = u"MB"
            mul   = value/1048576
        elif value < 1099511627775:
            width = 2
            unit  = u"GB"
            mul   = value/1073741824
        elif value < 1125899906842623:
            width = 2
            unit  = u"TB"
            mul   = value/1099511627776
        elif value < 1152921504606846975:
            width = 2
            unit  = u"PB"
            mul   = value/1125899906842624
        else:
            return (u"1", u">万PB", 1, 5)

        return (int(mul), unit, len(str(int(mul))), width)

    def _format_number_width(self, value):
        unit = self._format_number_unit(int(value))
        reserved = TX_EACH_WIDTH - (TX_NUMBER_WIDTH * unit[2] + TX_UNIT_DWORD_WIDTH * unit[3])
        return (unit[0], unit[1], reserved)

    def _format_traffic_width(self, value):
        unit = self._format_traffic_unit(int(value)*1024)
        reserved = TX_EACH_WIDTH - (TX_NUMBER_WIDTH * unit[2] + TX_UNIT_BASIC_WIDTH * unit[3])
        return (unit[0], unit[1], reserved)

    def _prepare_document(self):
        self._docx = os.path.join(self._path, str(uuid.uuid4()) + '.docx')
        shutil.copyfile(template_v1_h1_path, self._docx)
        self._doc = Document(self._docx)

        for i in range(0, P99_IDX_MAXIMUM_VALUE):
            self.__p_index_array.append(i)

        self.__dump_pindex_array()

    def _save_document(self, path):
        self._doc.save(path)
        os.remove(self._docx)

    def __dump_pindex_array(self):
        s = ''
        for i in self.__p_index_array:
            s += '{0:02}'.format(i) + ' '
        print s

    def __rearrange_pindex_array(self, position):
        self.__p_index_array[position] = -1

        for i in range(position + 1, P99_IDX_MAXIMUM_VALUE):
            self.__p_index_array[i] -= 1
        
        self.__dump_pindex_array()

    def _delete_paragraph(self, p_index):
        rp_index = self.__p_index_array[p_index]

        p = self._doc.paragraphs[rp_index]._element
        p.getparent().remove(p)
        p._p = p._element = None
        
        self.__rearrange_pindex_array(p_index)

    def _next_img(self, i_index = None):
        if i_index is not None:
            self._img_index = i_index
        else:
            self._img_index += 1

    def _skip_img(self, sha1_index = None):
        if sha1_index is not None:
            self._sha1_index = sha1_index
        else:
            self._sha1_index += 1

    def _get_img_sha1(self):
        return self._img_placeholders[self._sha1_index]

    def _get_img_name(self):
        return self._img_prefix + '{0:03}'.format(self._img_index) + self._img_suffix

    def _set_image(self, i_sha1, img_path):
        print 'i_sha1 -> ' + i_sha1 + ' img_path -> ' + img_path
        image_parts = self._doc.part.package.image_parts
        for p in image_parts:
            if p.sha1 == i_sha1:
                p._blob = open(img_path, 'r').read()
                break

    def set_image_in_order(self):
        self._set_image(self._get_img_sha1(), os.path.join(self._path, self._get_img_name()))
        self._next_img()
        self._skip_img()

    def _set_text(self, p_index, r_index, text):
        try:
            p = self._doc.paragraphs[self.__p_index_array[p_index]]
            r = p.runs[r_index]
            
            if not text:
                raise RuntimeError('NoneType has been detected')

            r.text = text
        except Exception, e:
            raise e

    def _set_int(self, p_index, r_index, number):
        self._set_text(p_index, r_index, str(number))
            
    def _set_float(self, p_index, r_index, number):
        self._set_text(p_index, r_index, '{0:.2f}'.format(number))

    def __text_alignment(self, text, t_width, quantity, qr_width, tl_width = 0):
        print 't_width -> ' + str(t_width) + ' qr_width -> ' + str(qr_width) + ' tl_width -> ' + str(tl_width)
        return ''.ljust(tl_width) + text + quantity + ''.ljust(int(t_width) - tl_width - qr_width)

    def _export_summary(self, data):

        if not data:
            self._delete_paragraph(P3_IDX_SUMMARY_CONTENT)
            self._delete_paragraph(P2_IDX_SUMMARY_TITLE)
        else:
            self._set_text(  P3_IDX_SUMMARY_CONTENT, R_IDX_P3_RANGE_TIME , self._range_time)

            self._set_int(P3_IDX_SUMMARY_CONTENT, R_IDX_P3_SUBDOMAIN     , data.get('domain_count'))
            self._set_int(P3_IDX_SUMMARY_CONTENT, R_IDX_P3_TOTAL_ATTACK  , data.get('attack_count'))
            self._set_int(P3_IDX_SUMMARY_CONTENT, R_IDX_P3_WEB_ATTACK    , data.get('web_attack_count'))
            self._set_int(P3_IDX_SUMMARY_CONTENT, R_IDX_P3_CC_ATTACK     , data.get('cc_count'))
            self._set_int(P3_IDX_SUMMARY_CONTENT, R_IDX_P3_IPADDRESS     , data.get('ip_count'))
            self._set_int(P3_IDX_SUMMARY_CONTENT, R_IDX_P3_TOTAL_REQUESTS, data.get('request_pv_total'))
            self._set_int(P3_IDX_SUMMARY_CONTENT, R_IDX_P3_ACC_TIMES     , data.get('hit_pv_total'))
            self._set_int(P3_IDX_SUMMARY_CONTENT, R_IDX_P3_TOTAL_TRAFFIC , data.get('request_flow_total'))
            self._set_int(P3_IDX_SUMMARY_CONTENT, R_IDX_P3_SAVED_TRAFFIC , data.get('hit_flow_total'))

            fmt_data = self._format_number_width(data.get('attack_count'))
            self._set_int( P4_IDX_STATISTICS_DATA, R_IDX_P4_T1_S1 , fmt_data[0])
            self._set_text(P4_IDX_STATISTICS_DATA, R_IDX_P4_T1_S2 , self.__text_alignment(fmt_data[1], fmt_data[2], u'次', 2))

            fmt_data = self._format_number_width(data.get('web_attack_count'))
            self._set_int( P4_IDX_STATISTICS_DATA, R_IDX_P4_T2_S1 , fmt_data[0])
            self._set_text(P4_IDX_STATISTICS_DATA, R_IDX_P4_T2_S2 , self.__text_alignment(fmt_data[1], fmt_data[2], u'次', 2))

            fmt_data = self._format_number_width(data.get('cc_count'))
            self._set_int( P4_IDX_STATISTICS_DATA, R_IDX_P4_T3_S1 , fmt_data[0])
            self._set_text(P4_IDX_STATISTICS_DATA, R_IDX_P4_T3_S2 , self.__text_alignment(fmt_data[1], fmt_data[2], u'次', 2))

            fmt_data = self._format_number_width(data.get('ip_count'))
            self._set_int( P4_IDX_STATISTICS_DATA, R_IDX_P4_T4_S1 , fmt_data[0])
            self._set_text(P4_IDX_STATISTICS_DATA, R_IDX_P4_T4_S2 , self.__text_alignment(fmt_data[1], fmt_data[2], u'个', 2))

            fmt_data = self._format_traffic_width(data.get('hit_flow_total'))
            self._set_int( P4_IDX_STATISTICS_DATA, R_IDX_P4_T5_S1 , fmt_data[0])
            self._set_text(P4_IDX_STATISTICS_DATA, R_IDX_P4_T5_S2 , self.__text_alignment(fmt_data[1], fmt_data[2], u'', 0, 2))

            fmt_data = self._format_number_width(data.get('hit_pv_total'))
            self._set_int( P4_IDX_STATISTICS_DATA, R_IDX_P4_T6_S1 , fmt_data[0])
            self._set_text(P4_IDX_STATISTICS_DATA, R_IDX_P4_T6_S2 , self.__text_alignment(fmt_data[1], fmt_data[2], u'次', 2))


    def _export_security(self, data):
    
        zone = (data and data.get('attack_tendency'))

        if not zone:
            self._delete_paragraph(P10_IDX_DEFENSE_ATTACK_TREND_PICTURE)
            self._delete_paragraph(P9_IDX_DEFENSE_ATTACK_TREND_CONTENT)
            self._delete_paragraph(P8_IDX_DEFENSE_ATTACK_TREND_TITLE)
            self._skip_img()
        else:
            self._set_text(P9_IDX_DEFENSE_ATTACK_TREND_CONTENT, R_IDX_P9_RANGE_TIME       , self._range_time)
            self._set_text(P9_IDX_DEFENSE_ATTACK_TREND_CONTENT, R_IDX_P9_LOW_TIME         , self.format_datetime(zone['min_count']['datetime']))
            self._set_int( P9_IDX_DEFENSE_ATTACK_TREND_CONTENT, R_IDX_P9_LOW_ATTACK_TIMES , zone['min_count']['attack_count'])
            self._set_text(P9_IDX_DEFENSE_ATTACK_TREND_CONTENT, R_IDX_P9_HIGH_TIME        , self.format_datetime(zone['max_count']['datetime']))
            self._set_int( P9_IDX_DEFENSE_ATTACK_TREND_CONTENT, R_IDX_P9_HIGH_ATTACK_TIMES, zone['max_count']['attack_count'])

            self.set_image_in_order()

        zone = (data and data.get('attack_domain'))

        if not zone:
            self._delete_paragraph(P13_IDX_DEFENSE_ATTACK_WEBSITE_PICTURE)
            self._delete_paragraph(P12_IDX_DEFENSE_ATTACK_WEBSITE_CONTENT)
            self._delete_paragraph(P11_IDX_DEFENSE_ATTACK_WEBSITE_TITLE)
            self._skip_img()
        else:
            
            self._set_text(P12_IDX_DEFENSE_ATTACK_WEBSITE_CONTENT, R_IDX_P12_RANGE_TIME , self._range_time)
            self._set_int(P12_IDX_DEFENSE_ATTACK_WEBSITE_CONTENT, R_IDX_P12_TOTAL_DOMAIN, zone['total'])
            
            sorted_items = sorted(zone['items'], key = lambda x: x['attack_count'], reverse = True)

            if len(sorted_items) == 0:
                self._delete_paragraph(P12_IDX_DEFENSE_ATTACK_WEBSITE_CONTENT)
            else:
                self._set_int(P12_IDX_DEFENSE_ATTACK_WEBSITE_CONTENT, R_IDX_P12_TOP_DOMAIN  , sorted_items[0]['attack_domain'])
                self._set_int(P12_IDX_DEFENSE_ATTACK_WEBSITE_CONTENT, R_IDX_P12_TOP_COUNT   , sorted_items[0]['attack_count'])
            
            self.set_image_in_order()

        zone = (data and data.get('attack_type'))

        if not zone:
            self._delete_paragraph(P16_IDX_DEFENSE_ATTACK_DISTRIBUTION_PICTURE)
            self._delete_paragraph(P15_IDX_DEFENSE_ATTACK_DISTRIBUTION_CONTENT)
            self._delete_paragraph(P14_IDX_DEFENSE_ATTACK_DISTRIBUTION_TITLE)
            self._skip_img()
        else:
            self._set_text(P15_IDX_DEFENSE_ATTACK_DISTRIBUTION_CONTENT, R_IDX_P15_RANGE_TIME  , self._range_time)
            
            sorted_items   = sorted(zone['items'], key = lambda x: x['attack_count'], reverse = True)

            if len(sorted_items) == 0:
                self._delete_paragraph(P15_IDX_DEFENSE_ATTACK_DISTRIBUTION_CONTENT)
            else:
                total_acttacks = sum([i['attack_count'] for i in sorted_items])
                percentage     = float(sorted_items[0]['attack_count']) / float(total_acttacks) * 100

                self._set_text( P15_IDX_DEFENSE_ATTACK_DISTRIBUTION_CONTENT, R_IDX_P15_ATTACK_CATEGORY, sorted_items[0]['attack_type'])
                self._set_int(  P15_IDX_DEFENSE_ATTACK_DISTRIBUTION_CONTENT, R_IDX_P15_ATTACK_TIMES   , int(sorted_items[0]['attack_count']))
                self._set_float(P15_IDX_DEFENSE_ATTACK_DISTRIBUTION_CONTENT, R_IDX_P15_PERCENTAGE     , percentage)

            self.set_image_in_order()

        zone = (data and data.get('attack_ip'))

        if not zone:
            self._delete_paragraph(P19_IDX_DEFENSE_SOURCE_IPADDRESS_PICTURE)
            self._delete_paragraph(P18_IDX_DEFENSE_SOURCE_IPADDRESS_CONTENT)
            self._delete_paragraph(P17_IDX_DEFENSE_SOURCE_IPADDRESS_TITLE)
            self._skip_img()
        else:
            self._set_text(P18_IDX_DEFENSE_SOURCE_IPADDRESS_CONTENT, R_IDX_P18_RANGE_TIME  , self._range_time)
            
            sorted_items   = sorted(zone['items'], key = lambda x: x['attack_count'], reverse = True)

            if len(sorted_items) == 0:
                self._delete_paragraph(P18_IDX_DEFENSE_SOURCE_IPADDRESS_CONTENT)
            else:
                total_acttacks = sum([i['attack_count'] for i in sorted_items])
                percentage     = float(sorted_items[0]['attack_count']) / float(total_acttacks) * 100

                self._set_text( P18_IDX_DEFENSE_SOURCE_IPADDRESS_CONTENT, R_IDX_P18_IP_ADDRESS  , sorted_items[0]['attack_ip'])
                self._set_int(  P18_IDX_DEFENSE_SOURCE_IPADDRESS_CONTENT, R_IDX_P18_ATTACK_TIMES, sorted_items[0]['attack_count'])

            self.set_image_in_order()

        zone = (data and data.get('attack_location'))

        if not zone:
            self._delete_paragraph(P22_IDX_DEFENSE_SOURCE_REGION_PICTURE)
            self._delete_paragraph(P21_IDX_DEFENSE_SOURCE_REGION_CONTENT)
            self._delete_paragraph(P20_IDX_DEFENSE_SOURCE_REGION_TITLE)
            self._skip_img()
        else:
            self._set_text(P21_IDX_DEFENSE_SOURCE_REGION_CONTENT, R_IDX_P21_RANGE_TIME   , self._range_time)
            
            sorted_items   = sorted(zone['items'], key = lambda x: x['attack_count'], reverse = True)

            if len(sorted_items) == 0:
                self._delete_paragraph(P21_IDX_DEFENSE_SOURCE_REGION_CONTENT)
            else:
                self._set_text(P21_IDX_DEFENSE_SOURCE_REGION_CONTENT, R_IDX_P21_REGION      , sorted_items[0]['attack_location'])
                self._set_int( P21_IDX_DEFENSE_SOURCE_REGION_CONTENT, R_IDX_P21_ATTACK_TIMES, sorted_items[0]['attack_count'])

            self.set_image_in_order()

        if not data:
            self._delete_paragraph(P7_IDX_DEFENSE_SUBJECT)

    def _export_accelerate(self, data):
   
        zone = (data and data.get('accel_request'))

        if not zone:
            self._delete_paragraph(P26_IDX_ACCELERATION_REQUESTS_PICTURE)
            self._delete_paragraph(P25_IDX_ACCELERATION_REQUESTS_CONTENT)
            self._delete_paragraph(P24_IDX_ACCELERATION_REQUESTS_TITLE)
            self._skip_img()
        else:
            self._set_text(P25_IDX_ACCELERATION_REQUESTS_CONTENT, R_IDX_P25_RANGE_TIME    , self._range_time)
            self._set_int( P25_IDX_ACCELERATION_REQUESTS_CONTENT, R_IDX_P25_TOTAL_REQUESTS, zone['total_request'])
            self._set_int( P25_IDX_ACCELERATION_REQUESTS_CONTENT, R_IDX_P25_ACC_REQUESTS  , zone['total_hit_pv'])

            self.set_image_in_order()

        zone = (data and data.get('accel_flow'))

        if not zone:
            self._delete_paragraph(P29_IDX_ACCELERATION_TRAFFIC_PICTURE)
            self._delete_paragraph(P28_IDX_ACCELERATION_TRAFFIC_CONTENT)
            self._delete_paragraph(P27_IDX_ACCELERATION_TRAFFIC_TITLE)
            self._skip_img()
        else:
            self._set_text(P28_IDX_ACCELERATION_TRAFFIC_CONTENT, R_IDX_P28_RANGE_TIME   , self._range_time)
            self._set_int( P28_IDX_ACCELERATION_TRAFFIC_CONTENT, R_IDX_P28_TOTAL_TRAFFIC, zone['total_flow'])
            self._set_int( P28_IDX_ACCELERATION_TRAFFIC_CONTENT, R_IDX_P28_SAVED_TRAFFIC, zone['total_hit_flow'])

            self.set_image_in_order()

        zone = (data and data.get('accel_ip'))

        if not zone:
            self._delete_paragraph(P32_IDX_ACCELERATION_IPADDRESS_PICTURE)
            self._delete_paragraph(P31_IDX_ACCELERATION_IPADDRESS_CONTENT)
            self._delete_paragraph(P30_IDX_ACCELERATION_IPADDRESS_TITLE)
            self._skip_img()
        else:
            self._set_text(P31_IDX_ACCELERATION_IPADDRESS_CONTENT, R_IDX_P31_RANGE_TIME, self._range_time)
            self._set_int( P31_IDX_ACCELERATION_IPADDRESS_CONTENT, R_IDX_P31_TOTAL_IP  , zone['total'])
            self._set_int( P31_IDX_ACCELERATION_IPADDRESS_CONTENT, R_IDX_P31_TOP_IP    , zone['max_count']['ip_count'])
            self._set_text(P31_IDX_ACCELERATION_IPADDRESS_CONTENT, R_IDX_P31_TOP_TIME  , self.format_datetime(zone['max_count']['datetime']))

            self.set_image_in_order()

        zone = (data and data.get('accel_pv'))

        if not zone:
            self._delete_paragraph(P35_IDX_ACCELERATION_PV_PICTURE)
            self._delete_paragraph(P34_IDX_ACCELERATION_PV_CONTENT)
            self._delete_paragraph(P33_IDX_ACCELERATION_PV_TITLE)
            self._skip_img()
        else:
            self._set_text(P34_IDX_ACCELERATION_PV_CONTENT, R_IDX_P34_RANGE_TIME, self._range_time)
            self._set_int( P34_IDX_ACCELERATION_PV_CONTENT, R_IDX_P34_TOTAL_PV  , zone['total'])
            self._set_int( P34_IDX_ACCELERATION_PV_CONTENT, R_IDX_P34_TOP_PV    , zone['max_count']['pv_count'])
            self._set_text(P34_IDX_ACCELERATION_PV_CONTENT, R_IDX_P34_TOP_TIME  , self.format_datetime(zone['max_count']['datetime']))

            self.set_image_in_order()

        zone = (data and data.get('accel_engines'))

        if not zone:
            self._delete_paragraph(P37_IDX_ACCELERATION_TREND_PICTURE)
            self._delete_paragraph(P36_IDX_ACCELERATION_TREND_TITLE)
            self._skip_img()
        else:
            self.set_image_in_order()

        zone = (data and data.get('accel_location'))

        if not zone:
            self._delete_paragraph(P39_IDX_ACCELERATION_REGION_PICTURE)
            self._delete_paragraph(P38_IDX_ACCELERATION_REGION_TITLE)
            self._skip_img()
        else:
            self.set_image_in_order()

        if not data:
            self._delete_paragraph(P23_IDX_ACCELERATION_SUBJECT)

    def set_image_prefix(self, prefix):
        self._img_prefix = prefix

    def set_img_suffix(self, suffix):
        self._img_suffix = suffix

    def export(self, docx_path):
        try:
            self._prepare_document()

            range_time = self.format_datetime(self._data.get('start_time'), True) + ' - ' + \
                         self.format_datetime(self._data.get('end_time'), True) + ' '

            self._range_time = range_time

            export_time = self.format_datetime(self._data.get('create_time'), True)

            self._set_text(P0_IDX_MAIN_TITLE   , R_IDX_P0_TITLE      , self._data.get('report_name'))
            self._set_text(P1_IDX_MAIN_SUBTITLE, R_IDX_P1_RANGE_TIME , range_time)
            self._set_text(P1_IDX_MAIN_SUBTITLE, R_IDX_P1_EXPORT_TIME, export_time)

            self._export_summary(self._data.get('summary'))
            self._export_security(self._data.get('security'))
            self._export_accelerate(self._data.get('accelerate'))
 
            self._save_document(docx_path)

            return True
        except Exception, e:
            print str(e)
            return False

    @staticmethod
    def dump_runs(document):
        p_index = 0
        for p in document.paragraphs:
            r_index = 0
            for r in p.runs:
                print 'p_index [' + str(p_index) + '] r_index [' + str(r_index) + '] -> ' + r.text
                r_index += 1
            print ''
            p_index += 1

    @staticmethod
    def dump_paragraphs(document):
        p_index = 0
        for p in document.paragraphs:
            print 'p_index [' + str(p_index) + '] -> ' + p.text
            p_index += 1

    @staticmethod
    def dump_inline_shapes(document):
        s_index = 0
        for s in document.part.inline_shapes:
            shape_type = 'PIC (' if s.type == 3 else 'UNKNOWN'
            print 's_index [' + str(s_index) + '] -> ' + shape_type + str(s.height.cm) + ',' + str(s.width.cm) + ')'
            s_index += 1

    @staticmethod
    def dump_images(document):
        image_parts = document.part.package.image_parts
        part_index = 0
        for p in image_parts:
            print 'part_index [' + str(part_index) + '] -> ' + p.filename + ' sha1 -> ' + p.sha1
            part_index += 1


class PDFFormat(object):

    def __init__(self, path, host = '127.0.0.1'):
        self._path   = path
        self._prefix = 'report-'
        self._suffix = '.pdf'
        self._target = None
        self.__host  = host

    def export(self, report_id, start_time, end_time, cycle):

        self._target     = os.path.join(self._path, self._prefix + str(report_id) + self._suffix)

        wkhtmltopdf_url  = '"http://' + self.__host + '/views/report.html#/?report_id=%s&start_time=%s&end_time=%s&cycle=%s" ' % (report_id, start_time, end_time, cycle)
        
        wkhtmltopdf_args = '/usr/bin/wkhtmltopdf -q --orientation Landscape --javascript-delay 1000 --no-stop-slow-scripts --debug-javascript '

        wkhtmltopdf_args = wkhtmltopdf_args + wkhtmltopdf_url
        wkhtmltopdf_args = wkhtmltopdf_args + self._target

        try:
            exit_code = subprocess32.call(wkhtmltopdf_args, shell = True, timeout=20)
            if exit_code != 0 or not os.path.exists(self._target):
                return False
            else:
                return True
        except Exception, e:
            return False

    def extract_image(self):

        if not os.path.exists(self._target):
            return False

        pdfimages_args = '/usr/bin/pdfimages -j ' + self._target + ' ' + os.path.join(self._path, 'report')

        try:
            exit_code = subprocess32.call(pdfimages_args, shell = True, timeout=5)
            if exit_code != 0:
                return False
            else:
                return True
        except Exception, e:
            return False

class ZipReport(object):

    def __init__(self, report_id, file_name, start_time, end_time, cycle, host = '127.0.0.1'):
        self._cycle       = cycle
        self._report_id   = report_id
        self._file_name = file_name
        self._start_time  = start_time
        self._end_time    = end_time
        self.__temp_path  = os.path.join('/tmp', str(uuid.uuid4()))
        self.__host       = host
        self.__zip_report = None

    def export(self):
        try:

            if not os.path.exists(self.__temp_path):
                os.makedirs(self.__temp_path)

            pdf_format = PDFFormat(self.__temp_path, self.__host)

            success = pdf_format.export(self._report_id, self._start_time, self._end_time, self._cycle)
            if not success:
                raise RuntimeError('Export PDF format error')

            success = pdf_format.extract_image()
            if not success:
                raise RuntimeError('Extract images from PDF format error')

            data_url = 'http://' + self.__host + '/api/dashboard/report/%s?end_time=%s&start_time=%s' % (self._report_id, self._end_time, self._start_time)
            r = requests.get(data_url)
            if r.status_code != 200:
                raise RuntimeError('Fail to request data from server')

            docx_format = DocxFormat(r.json(), embedded_placeholder_sha1, self.__temp_path)
            success = docx_format.export(os.path.join(self.__temp_path, 'report-' + str(self._report_id) + '.docx'))
            if not success:
                raise RuntimeError('Export DOCX format error')

            pdf_name     = self._file_name + '.pdf'
            docx_name    = self._file_name + '.docx'
            self.__zip_report = os.path.join(self.__temp_path, 'report-' + str(self._report_id) + '.zip')

            zfd = zipfile.ZipFile(self.__zip_report, 'w')
            zfd.write(os.path.join(self.__temp_path, 'report-' + str(self._report_id) + '.pdf'), pdf_name)
            zfd.write(os.path.join(self.__temp_path, 'report-' + str(self._report_id) + '.docx'), docx_name)
            zfd.close()

            return True
        except Exception, e:
            print str(e)
            shutil.rmtree(self.__temp_path)
            return False

    def get_zip_report(self):
        return self.__zip_report

    def get_temp_path(self):
        return self.__temp_path


class ZipWordReport(object):

    def __init__(self, report_id, file_name, start_time, end_time, cycle, host = '127.0.0.1'):
        self._cycle       = cycle
        self._report_id   = report_id
        self._start_time  = start_time
        self._file_name = file_name
        self._end_time    = end_time
        self.__temp_path  = os.path.join('/tmp', str(uuid.uuid4()))
        self.__host       = host
        self.__zip_report = None

    def export(self):
        try:

            if not os.path.exists(self.__temp_path):
                os.makedirs(self.__temp_path)

            pdf_format = PDFFormat(self.__temp_path, self.__host)

            success = pdf_format.export(self._report_id, self._start_time, self._end_time, self._cycle)
            if not success:
                raise RuntimeError('Export PDF format error')

            success = pdf_format.extract_image()
            if not success:
                raise RuntimeError('Extract images from PDF format error')

            data_url = 'http://' + self.__host + '/api/dashboard/report/%s?end_time=%s&start_time=%s' % (self._report_id, self._end_time, self._start_time)
            r = requests.get(data_url)
            if r.status_code != 200:
                raise RuntimeError('Fail to request data from server')

            docx_format = DocxFormat(r.json(), embedded_placeholder_sha1, self.__temp_path)
            success = docx_format.export(os.path.join(self.__temp_path, 'report-' + str(self._report_id) + '.docx'))
            if not success:
                raise RuntimeError('Export DOCX format error')

            docx_name    = self._file_name  + '.docx'
            self.__zip_report = os.path.join(self.__temp_path, 'report-' + str(self._report_id) + '.zip')

            zfd = zipfile.ZipFile(self.__zip_report, 'w')
            zfd.write(os.path.join(self.__temp_path,  'report-' + str(self._report_id) + '.docx'), docx_name)
            zfd.close()

            return True
        except Exception, e:
            print str(e)
            shutil.rmtree(self.__temp_path)
            return False

    def get_zip_report(self):
        return self.__zip_report

    def get_temp_path(self):
        return self.__temp_path


class ZipPDFReport(object):

    def __init__(self, report_id, file_name, start_time, end_time, cycle, host = '127.0.0.1'):
        self._cycle       = cycle
        self._report_id   = report_id
        self._file_name = file_name
        self._start_time  = start_time
        self._end_time    = end_time
        self.__temp_path  = os.path.join('/tmp', str(uuid.uuid4()))
        self.__host       = host
        self.__zip_report = None

    def export(self):
        try:

            if not os.path.exists(self.__temp_path):
                os.makedirs(self.__temp_path)

            pdf_format = PDFFormat(self.__temp_path, self.__host)

            success = pdf_format.export(self._report_id, self._start_time, self._end_time, self._cycle)
            if not success:
                raise RuntimeError('Export PDF format error')

            pdf_name     = self._file_name + '.pdf'
            self.__zip_report = os.path.join(self.__temp_path, 'report-' + str(self._report_id) + '.zip')

            zfd = zipfile.ZipFile(self.__zip_report, 'w')
            zfd.write(os.path.join(self.__temp_path, 'report-' + str(self._report_id) + '.pdf'), pdf_name)
            zfd.close()

            return True
        except Exception, e:
            print str(e)
            shutil.rmtree(self.__temp_path)
            return False

    def get_zip_report(self):
        return self.__zip_report

    def get_temp_path(self):
        return self.__temp_path


# dump tempates structure information to make development ease
if __name__ == '__main__':

    # document = Document(template_v1_h1_path)
    # DocxFormat.dump_paragraphs(document)
    # DocxFormat.dump_runs(document)
    # DocxFormat.dump_inline_shapes(document)
    # DocxFormat.dump_images(document)

    with open('json-data.txt', 'r') as fd:
        data = json.load(fd)

    docx_format = DocxFormat(data, embedded_placeholder_sha1, workspace_home)
    docx_format.export(os.path.join(workspace_home, 'report.docx'))